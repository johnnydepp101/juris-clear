from fpdf import FPDF
from docx import Document
from io import BytesIO
import os
import streamlit as st

def create_pdf(text):
    pdf = FPDF()
    pdf.set_auto_page_break(auto=True, margin=15) # Авто-перенос страниц
    pdf.add_page()
    
    # Путь к шрифту
    font_path = "DejaVuSans.ttf" 
    
    # Новая логика проверки шрифта
    if not os.path.exists(font_path):
        st.error("Критическая ошибка: Файл шрифта DejaVuSans.ttf не найден. PDF не может быть создан.")
        return None
        
    pdf.add_font('DejaVu', '', font_path)
    pdf.set_font('DejaVu', '', 12)
    
    # Используем write вместо multi_cell. 
    # 10 — это высота строки. write() сам переносит текст и слова.
    pdf.write(10, text.strip())
    
    return pdf.output()

def create_docx(text):
    doc = Document()
    doc.add_heading('Результат анализа договора - JurisClear AI', 0)
    
    # Разбиваем текст на параграфы для красоты
    for paragraph in text.strip().split('\n'):
        if paragraph.strip():
            doc.add_paragraph(paragraph)
    
    bio = BytesIO()
    doc.save(bio)
    return bio.getvalue()
