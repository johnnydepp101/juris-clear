import streamlit as st
from openai import OpenAI
import pdfplumber
import re
from supabase import create_client, Client  # Добавили импорт Supabase
import os
from fpdf import FPDF
from docx import Document
from io import BytesIO
import pytesseract
from pdf2image import convert_from_bytes
from PIL import Image
import json
from streamlit.components.v1 import html

# --- 1. НАСТРОЙКА СТРАНИЦЫ ---
st.set_page_config(
    page_title="JurisClear AI",
    page_icon="⚖️",
    layout="wide",
    initial_sidebar_state="collapsed"
)

# --- ИНИЦИАЛИЗАЦИЯ ---
if 'reset_counter' not in st.session_state:
    st.session_state.reset_counter = 0
if 'user' not in st.session_state:
    st.session_state.user = None

# --- ВОССТАНОВЛЕНИЕ СЕССИИ ИЗ URL (ДЛЯ JS МОСТА) ---
if st.session_state.user is None:
    params = st.query_params
    if "access_token" in params and "refresh_token" in params:
        try:
            # Используем supabase_auth, так как это Auth-операция
            st.toast("🔄 Восстановление сессии...")
            session = supabase_auth.auth.set_session(params["access_token"], params["refresh_token"])
            st.session_state.user = session.user
            # Сохраняем в session_state для JS моста (чтобы он не удалил из localStorage)
            st.session_state.session_data = {
                "access_token": params["access_token"],
                "refresh_token": params["refresh_token"]
            }
            # Очищаем параметры URL
            st.query_params.clear()
            st.rerun()
        except Exception as e:
            st.toast(f"❌ Ошибка: {e}")
            pass

# --- 2. ВЕСЬ ДИЗАЙН (ПРЕМИАЛЬНЫЙ АДАПТИВНЫЙ CSS) ---
st.markdown("""
<style>
@import url('https://fonts.googleapis.com/css2?family=Inter:wght@400;500;600;700;800&display=swap');
html, body, [data-testid="stAppViewContainer"] {
    font-family: 'Inter', sans-serif !important;
}
:root {
    --bg-color: #0c0e12;
    --card-bg: rgba(22, 27, 34, 0.7);
    --text-color: #e2e8f0;
    --secondary-text: #94a3b8;
    --border-color: rgba(255, 255, 255, 0.08); /* Ультра-тонкие границы */
    --accent-blue: #3266e3; /* Royal Blue */
    --accent-cyan: #4cc9f0;
    --accent-green: #2ecc71;
    --glass-blur: blur(8px); /* Более чистое размытие */
    --card-shadow: 0 10px 40px rgba(0, 0, 0, 0.4);
    --header-color: #ffffff;
    --glow-color: rgba(50, 102, 227, 0.2);
}
@media (prefers-color-scheme: light) {
    :root {
        --bg-color: #f8fafc;
        --card-bg: rgba(255, 255, 255, 0.8);
        --text-color: #0f172a;
        --secondary-text: #64748b;
        --border-color: rgba(0, 0, 0, 0.05);
        --card-shadow: 0 10px 40px rgba(0, 0, 0, 0.05);
        --header-color: #1e293b;
        --glow-color: rgba(50, 102, 227, 0.05);
    }
}
#MainMenu {visibility: hidden;} footer {visibility: hidden;} header {visibility: hidden;}
[data-testid="stHeader"] {display: none;}
.block-container {
    padding-top: 3rem; /* Больше воздуха сверху */
    max-width: 100%;
    padding-left: 6rem;
    padding-right: 6rem;
}
@media (max-width: 1200px) {
    .block-container {
        padding-left: 3rem;
        padding-right: 3rem;
    }
}
@media (max-width: 768px) {
    .block-container {
        padding-left: 1.5rem;
        padding-right: 1.5rem;
        padding-top: 2rem;
    }
    h1 {
        font-size: 26px !important;
    }
    .stMarkdown div[data-testid="stMarkdownHeader"] h1 {
        font-size: 26px !important;
    }
}
[data-testid="stAppViewContainer"] {
    background-color: var(--bg-color);
    color: var(--text-color);
    transition: background-color 0.5s ease;
}
.stMarkdown h1 a, .stMarkdown h2 a, .stMarkdown h3 a, 
.stMarkdown h4 a, .stMarkdown h5 a, .stMarkdown h6 a { display: none !important; }
[data-testid="stMarkdownHeader"] a { display: none !important; }
[data-testid="stHorizontalBlock"] { align-items: center !important; }
.pricing-card-container {
    display: flex;
    gap: 20px;
    margin-bottom: 2rem;
    flex-wrap: wrap;
}
.pricing-card {
    flex: 1;
    min-width: 300px;
    background: var(--card-bg);
    backdrop-filter: var(--glass-blur);
    -webkit-backdrop-filter: var(--glass-blur);
    padding: 40px; 
    border-radius: 20px; 
    border: 1px solid var(--border-color); 
    text-align: left; 
    color: var(--text-color);
    box-shadow: var(--card-shadow);
    transition: transform 0.4s cubic-bezier(0.25, 1, 0.5, 1), box-shadow 0.4s ease, border-color 0.4s ease;
    position: relative;
    overflow: hidden;
}
.pricing-card:hover {
    transform: translateY(-8px);
    border-color: rgba(50, 102, 227, 0.4);
    box-shadow: 0 20px 60px rgba(0, 0, 0, 0.5);
}
.pricing-card::before {
    content: "";
    position: absolute;
    top: 0; left: 0; width: 100%; height: 2px;
    background: var(--accent-blue);
    opacity: 0.5;
}
.pricing-card-pro::before {
    background: var(--accent-green);
}
.pricing-header {
    font-size: 1.2rem;
    font-weight: 600;
    margin-bottom: 0.5rem;
    color: var(--secondary-text);
    text-transform: uppercase;
    letter-spacing: 1px;
}
.pricing-price {
    font-size: 2.5rem;
    font-weight: 800;
    margin-bottom: 1.5rem;
    color: var(--header-color);
}
.pricing-features {
    list-style: none;
    padding: 0;
    margin-bottom: 2rem;
}
.pricing-features li {
    margin-bottom: 12px;
    display: flex;
    align-items: center;
    gap: 10px;
    font-size: 0.95rem;
}
.pricing-features li::before {
    content: "✦";
    color: var(--accent-blue);
    font-size: 0.8rem;
}
.report-card {
    background-color: var(--card-bg);
    backdrop-filter: var(--glass-blur);
    -webkit-backdrop-filter: var(--glass-blur);
    border: 1px solid var(--border-color);
    padding: 50px; 
    border-radius: 20px; 
    margin-top: 30px; 
    color: var(--text-color);
    box-shadow: var(--card-shadow);
}
@keyframes fadeIn {
    from { opacity: 0; transform: translateY(20px); }
    to { opacity: 1; transform: translateY(0); }
}
.stAppViewContainer [data-testid="stVerticalBlock"] > div {
    animation: fadeIn 0.6s ease-out forwards;
}
.risk-meter-container {
    background: rgba(0, 0, 0, 0.2); 
    border-radius: 12px; 
    padding: 4px;
    border: 1px solid var(--border-color); 
    margin: 30px 0;
    overflow: hidden;
}
.risk-meter-bar {
    height: 30px; 
    border-radius: 8px; 
    display: flex; 
    align-items: center; 
    justify-content: center; 
    color: white; 
    font-weight: 700; 
    font-size: 14px;
    letter-spacing: 0.5px;
    text-shadow: 0 1px 2px rgba(0,0,0,0.2);
    transition: width 1s ease-in-out;
}
.stButton > button, .stLinkButton > a, .stDownloadButton > button {
    border-radius: 12px !important;
    font-weight: 600 !important;
    letter-spacing: 0.3px !important;
    font-size: 14px !important;
    padding: 0.75rem 2rem !important;
    transition: all 0.3s ease !important;
    border: 1px solid var(--border-color) !important;
    background: rgba(255, 255, 255, 0.03) !important;
    color: var(--text-color) !important;
}
.stButton > button:hover {
    border-color: var(--accent-blue) !important;
    background: rgba(50, 102, 227, 0.1) !important;
    transform: translateY(-2px);
}
.stButton > button[kind="primary"] {
    background: var(--accent-blue) !important;
    color: white !important;
    border: none !important;
}
[data-testid="stFileUploader"] {
    background: var(--card-bg);
    padding: 3rem;
    border-radius: 20px;
    border: 1px dashed var(--border-color);
    transition: background 0.3s ease, border-color 0.3s ease;
}
[data-testid="stFileUploader"]:hover {
    border-color: var(--accent-blue);
    background: rgba(50, 102, 227, 0.03);
}
.secondary-text {
    color: var(--secondary-text);
    font-size: 0.9rem;
}
</style>
""", unsafe_allow_html=True)

# --- JS МОСТ ДЛЯ LOCALSTORAGE ---
# Этот скрипт синхронизирует сессию Supabase между браузером и Streamlit
user_status = "logged_in" if st.session_state.user else "logged_out"
current_session = "{}"
if st.session_state.user and 'session_data' in st.session_state:
    current_session = json.dumps(st.session_state.session_data)

st.components.v1.html(f"""
<script>
    const US_KEY = 'supabase.auth.token';
    const status = "{user_status}";
    const sessionToSave = {current_session};

    console.log("[JurisClear Bridge] Python Status:", status);

    // 1. Если пользователь вошел в Python, но в localStorage пусто - сохраняем
    if (status === "logged_in" && Object.keys(sessionToSave).length > 0) {{
        console.log("[JurisClear Bridge] Saving session to localStorage...");
        localStorage.setItem(US_KEY, JSON.stringify(sessionToSave));
    }}

    // 2. Если пользователь вышел в Python, но в localStorage еще есть данные - удаляем
    if (status === "logged_out") {{
        if (localStorage.getItem(US_KEY)) {{
            console.log("[JurisClear Bridge] Clearing session from localStorage...");
            localStorage.removeItem(US_KEY);
        }}
    }}

    // 3. Главная магия: Авто-восстановление при загрузке
    const savedSessionStr = localStorage.getItem(US_KEY);
    const urlParams = new URLSearchParams(window.parent.location.search);
    
    if (status === "logged_out" && savedSessionStr && !urlParams.has('access_token')) {{
        try {{
            const savedSession = JSON.parse(savedSessionStr);
            if (savedSession && savedSession.access_token && savedSession.refresh_token) {{
                console.log("[JurisClear Bridge] Session found in localStorage! Redirecting parent window...");
                urlParams.set('access_token', savedSession.access_token);
                urlParams.set('refresh_token', savedSession.refresh_token);
                window.parent.location.search = urlParams.toString();
            }}
        }} catch (e) {{
            console.error("[JurisClear Bridge] Error parsing session:", e);
        }}
    }} else {{
        console.log("[JurisClear Bridge] No session restoration needed.");
    }}
</script>
""", height=0)

# --- 3. ЛОГИКА ДИНАМИЧЕСКОЙ ШКАЛЫ ---
def get_risk_params(score):
    if score <= 3:
        return "linear-gradient(90deg, #059669 0%, #10b981 100%)", "rgba(16, 185, 129, 0.5)", "НИЗКИЙ"
    elif score <= 6:
        return "linear-gradient(90deg, #d97706 0%, #fbbf24 100%)", "rgba(251, 191, 36, 0.5)", "СРЕДНИЙ"
    else:
        return "linear-gradient(90deg, #dc2626 0%, #ef4444 100%)", "rgba(239, 68, 68, 0.5)", "КРИТИЧЕСКИЙ"

# --- 4. ПОДКЛЮЧЕНИЕ API И БАЗЫ ДАННЫХ ---
# OpenAI
client = OpenAI(api_key=st.secrets["OPENAI_API_KEY"])

# Supabase
# Попытка инициализации из секретов Streamlit (рекомендуемый способ)
try:
    url: str = st.secrets["SUPABASE_URL"]
    # Пробуем взять Service Key для полной свободы действий, если он есть, иначе Anon Key
    key: str = st.secrets.get("SUPABASE_SERVICE_KEY") or st.secrets.get("SUPABASE_KEY")
    supabase: Client = create_client(url, key)
    # Auth-клиент (anon key) — для регистрации/входа (service key НЕ подходит для auth)
    supabase_auth: Client = create_client(url, st.secrets["SUPABASE_KEY"])
except Exception as e:
    st.error(f"Ошибка подключения к Supabase: {e}. Проверьте secrets.toml")
    # Создаем фиктивный клиент, чтобы приложение не вылетало сразу при отрисовке UI
    supabase = None
    supabase_auth = None

# --- ФУНКЦИЯ ИЗВЛЕЧЕНИЯ ТЕКСТА (ОБНОВЛЕННАЯ С ГИБРИДНЫМ OCR) ---
def extract_text_from_pdf(file_bytes):
    """
    Гибридная функция: сначала пробует вытащить текст напрямую, 
    если не выходит — использует OCR (Tesseract).
    """
    text = ""
    # 1. Пробуем стандартный метод (pdfplumber)
    try:
        with pdfplumber.open(BytesIO(file_bytes)) as pdf:
            for page in pdf.pages:
                content = page.extract_text()
                if content:
                    text += content + "\n"
    except Exception as e:
        st.warning(f"Обычное чтение PDF не удалось, пробуем OCR... ({e})")

    # 2. Если текста почти нет (< 100 символов), значит это скан или фото
    if len(text.strip()) < 100:
        with st.status("🔍 Обнаружен скан. Работает OCR (распознавание текста)...", expanded=True) as status:
            try:
                # Конвертируем PDF в список изображений (каждая страница = картинка)
                images = convert_from_bytes(file_bytes)
                
                ocr_text = ""
                total_pages = len(images)
                
                for i, image in enumerate(images):
                    st.write(f"Обработка страницы {i+1} из {total_pages}...")
                    # Распознаем текст (русский + английский)
                    page_text = pytesseract.image_to_string(image, lang='rus+eng')
                    ocr_text += f"\n--- Страница {i+1} ---\n" + page_text
                
                status.update(label="✅ Распознавание завершено!", state="complete", expanded=False)
                return ocr_text
            except Exception as e:
                st.error(f"Ошибка OCR: {e}. Убедитесь, что tesseract-ocr прописан в packages.txt")
                return ""
    
    return text

# --- ФУНКЦИЯ СОЗДАНИЯ PDF (ИНТЕГРИРОВАНА НОВАЯ ПРОВЕРКА ШРИФТА) ---
def create_pdf(text):
    pdf = FPDF()
    pdf.set_auto_page_break(auto=True, margin=15) # Авто-перенос страниц
    pdf.add_page()
    
    # Путь к шрифту
    font_path = "DejaVuSans.ttf" 
    
    # Новая логика проверки шрифта
    if not os.path.exists(font_path):
        st.error("Критическая ошибка: Файл шрифта DejaVuSans.ttf не найден. PDF не может быть создан.")
        return None
        
    pdf.add_font('DejaVu', '', font_path)
    pdf.set_font('DejaVu', '', 12)
    
    # Используем write вместо multi_cell. 
    # 10 — это высота строки. write() сам переносит текст и слова.
    pdf.write(10, text.strip())
    
    return pdf.output()

# --- ФУНКЦИЯ СОЗДАНИЯ WORD ---
def create_docx(text):
    doc = Document()
    doc.add_heading('Результат анализа договора - JurisClear AI', 0)
    
    # Разбиваем текст на параграфы для красоты
    for paragraph in text.strip().split('\n'):
        if paragraph.strip():
            doc.add_paragraph(paragraph)
    
    bio = BytesIO()
    doc.save(bio)
    return bio.getvalue()

# --- ФУНКЦИЯ ДЛЯ АНАЛИЗА ДЛИННЫХ ТЕКСТОВ ---
def analyze_long_text(full_text, contract_type, user_role, special_instructions, prompt_instruction):
    # Делим текст по абзацам, а не по символам
    paragraphs = full_text.split('\n\n')
    chunks = []
    current_chunk = ""
    
    for p in paragraphs:
        if len(current_chunk) + len(p) < 15000:
            current_chunk += p + "\n\n"
        else:
            chunks.append(current_chunk)
            current_chunk = p + "\n\n"
    if current_chunk:
        chunks.append(current_chunk)
    
    partial_analyses = []
    
    # Индикатор прогресса для пользователя
    progress_bar = st.progress(0)
    status_text = st.empty()
    
    for idx, chunk in enumerate(chunks):
        status_text.text(f"Анализирую часть {idx+1} из {len(chunks)}...")
        
        prompt = (
            f"Ты — опытный юрист. Проанализируй эту ЧАСТЬ договора ({idx+1}/{len(chunks)}). "
            f"Тип: {contract_type}, Роль: {user_role}. {special_instructions}\n\n"
            "Выдели только КРИТИЧЕСКИЕ риски и кабальные условия, которые видишь в этом куске.\n"
            f"Текст части:\n{chunk}"
        )
        
        response = client.chat.completions.create(
            model="gpt-4o-mini", # Используем mini для промежуточных частей (быстрее и дешевле)
            messages=[{"role": "user", "content": prompt}],
            temperature=0.0
        )
        partial_analyses.append(response.choices[0].message.content)
        progress_bar.progress((idx + 1) / len(chunks))

    # 2. Финальная сборка отчета
    status_text.text("Формирую итоговый отчет и протокол разногласий...")
    
    combined_context = "\n\n".join(partial_analyses)
    
    final_prompt = (
        "Перед тобой промежуточные результаты анализа разных частей одного договора. "
        "Твоя задача — объединить их в один профессиональный, структурированный отчет.\n\n"
        "УДАЛИ повторы. СГРУППИРУЙ риски по категориям (Финансовые, Сроки, Ответственность).\n"
        "СОСТАВЬ итоговую таблицу 'Протокол разногласий'.\n\n"
        "СТРУКТУРА И ЯЗЫК ОТВЕТА ДОЛЖНЫ СТРОГО СООТВЕТСТВОВАТЬ ЭТАЛОНУ (с SCORE).\n\n"
        f"Промежуточные данные:\n{combined_context}"
    )
    
    final_response = client.chat.completions.create(
        model="gpt-4o", # Для финала используем мощную модель
        messages=[
            {"role": "system", "content": prompt_instruction}, # Используем твой основной системный промпт
            {"role": "user", "content": final_prompt}
        ],
        temperature=0.0
    )
    
    status_text.empty()
    progress_bar.empty()
    
    return final_response.choices[0].message.content

# --- НОВАЯ ФУНКЦИЯ АНАЛИЗА (INTEGRATED) ---
def generate_analysis(full_text, contract_type, user_role, special_reqs):

    # --- НАСТРОЙКИ ЧАНКИНГА ---
    MAX_CHUNK_SIZE = 15000  # Примерно 3000-4000 токенов

    # 1. Если текст короткий, анализируем в один проход
    if len(full_text) < MAX_CHUNK_SIZE:
        prompt = f"""
        Ты — профессиональный корпоративный юрист. Проведи полный аудит договора: {contract_type}.
        Моя роль: {user_role}. Особые требования: {special_reqs if special_reqs else 'Нет'}.
        Текст договора:
        {full_text}
        
        Выдай отчет в Markdown: Резюме, Риски (Критический/Средний/Низкий), Топ-3 опасных пункта, Рекомендации.
        """
        response = client.chat.completions.create(
            model="gpt-4o",
            messages=[{"role": "user", "content": prompt}],
            temperature=0.2
        )
        return response.choices[0].message.content

    # 2. Если текст длинный — включаем логику ЧАНКИНГА
    st.info(f"⏳ Договор очень длинный ({len(full_text)} симв.). Анализирую по частям...")
    
    paragraphs = full_text.split('\n\n')
    chunks = []
    current_chunk = ""
    for p in paragraphs:
        if len(current_chunk) + len(p) < MAX_CHUNK_SIZE:
            current_chunk += p + "\n\n"
        else:
            chunks.append(current_chunk)
            current_chunk = p + "\n\n"
    if current_chunk:
        chunks.append(current_chunk)

    # Собираем риски из каждой части
    partial_risks = []
    progress_bar = st.progress(0)
    
    for i, chunk in enumerate(chunks):
        st.write(f"Сканирую часть {i+1} из {len(chunks)}...")
        chunk_prompt = f"Найди все юридические риски в этой части договора {contract_type} для роли {user_role}. Текст:\n{chunk}"
        
        chunk_res = client.chat.completions.create(
            model="gpt-4o",
            messages=[{"role": "user", "content": chunk_prompt}],
            temperature=0
        )
        partial_risks.append(chunk_res.choices[0].message.content)
        progress_bar.progress((i + 1) / len(chunks))

    # --- ФИНАЛЬНЫЙ СИНТЕЗ ---
    st.write("Сборка финального отчета...")
    all_risks_text = "\n\n".join(partial_risks)
    
    final_prompt = f"""
    Ты — старший юрист. Перед тобой список рисков, найденных в разных частях одного договора ({contract_type}).
    Твоя задача: объединить их в один логичный, структурированный отчет для клиента ({user_role}).
    Исключи повторы и выдели самые критические моменты.
    
    Список всех найденных рисков:
    {all_risks_text}
    
    Выдай финальный отчет в Markdown: Резюме, Общая оценка рисков, Детальный список ловушек, Рекомендации.
    """
    
    final_res = client.chat.completions.create(
        model="gpt-4o",
        messages=[{"role": "user", "content": final_prompt}],
        temperature=0.2
    )
    return final_res.choices[0].message.content

# === НОВЫЙ ПРОФЕССИОНАЛЬНЫЙ ПРИМЕР ОТЧЕТА ===
sample_text = """
### 📋 КРАТКОЕ РЕЗЮМЕ АУДИТА: ДОГОВОР ОКАЗАНИЯ УСЛУГ

**ОБЩИЙ ВЕРДИКТ:** Договор составлен с существенным перекосом баланса интересов в пользу Исполнителя и содержит условия, способные нанести значительный финансовый ущерб Заказчику. Настоятельно рекомендуется доработка перед подписанием.

---

#### 1. ФИНАНСОВЫЕ РИСКИ И ОТВЕТСТВЕННОСТЬ

**🔴 Критическая угроза: Кабальная неустойка (Пункт 6.1)**
* **Суть условия:** Установлена пеня за просрочку оплаты в размере **1% в день** от суммы задолженности.
* **Юридический анализ:** Это эквивалентно **365% годовых**, что более чем в 10 раз превышает стандартную деловую практику (обычно 0,1%). Суд с высокой вероятностью признает такую неустойку несоразмерной, но до этого момента вы будете накапливать огромный долг. Риск потери ликвидности.

**🟠 Высокая угроза: Неконтролируемое изменение цены (Пункт 4.2)**
* **Суть условия:** Исполнитель имеет право в одностороннем порядке повышать стоимость услуг, уведомив Заказчика за 5 рабочих дней.
* **Юридический анализ:** Отсутствует механизм согласования новой цены или безусловное право Заказчика на расторжение договора без штрафов в случае несогласия с новой ценой. Риск непланируемых расходов.

#### 2. РИСКИ РАСТОРЖЕНИЯ И РАЗРЕШЕНИЯ СПОРОВ

**🟡 Средняя угроза: Невыгодная договорная подсудность (Пункт 9.3)**
* **Суть условия:** Все споры по договору подлежат рассмотрению в арбитражном суде по месту нахождения Исполнителя (г. Владивосток).
* **Юридический анализ:** Это существенно усложняет и удорожает процесс защиты ваших прав (транспортные расходы, наем локальных представителей), если ваша компания находится в другом регионе.

*💡 (Примечание: Полная версия отчета содержит конкретные формулировки правок (протокол разногласий) для нейтрализации каждого из указанных рисков.)*
"""



# --- 6. ИНТЕРФЕЙС ПРИЛОЖЕНИЯ (доступен всем) ---

# --- ХЕДЕР ПРИЛОЖЕНИЯ ---
header_col1, header_col2 = st.columns([3, 1])

with header_col1:
    st.markdown(f"""
<div style="display: flex; align-items: center; gap: 15px;">
<span style="font-size: 40px; line-height: 1;">⚖️</span>
<div style="display: flex; flex-direction: column;">
<h1 style='color: var(--header-color); margin: 0; padding: 0; font-size: 32px; font-weight: 800; line-height: 1;'>JurisClear <span style='color:var(--accent-blue)'>AI</span></h1>
</div>
</div>
""", unsafe_allow_html=True)

with header_col2:
    if st.session_state.user:
        user_email = st.session_state.user.email
        with st.popover(f"👤 {user_email}", use_container_width=True):
            st.markdown(f"""
<div style="padding: 5px 0;">
<p style="margin: 0 0 5px 0; font-size: 13px; color: var(--secondary-text);">Вы вошли как:</p>
<p style="margin: 0 0 15px 0; font-weight: 700; color: var(--text-color); font-size: 14px;">{user_email}</p>
</div>
""", unsafe_allow_html=True)
            if st.button("🚪 Выйти из аккаунта", use_container_width=True, key="btn_logout"):
                try:
                    supabase_auth.auth.sign_out()
                except Exception:
                    pass
                st.session_state.user = None
                keys_to_clear = ["analysis_result", "audit_score", "session_data"]
                for k in keys_to_clear:
                    if k in st.session_state:
                        del st.session_state[k]
                st.rerun()
    else:
        with st.popover("🔐 Войти", use_container_width=True):
            tab_login, tab_register = st.tabs(["🔑 Вход", "📝 Регистрация"])

            with tab_login:
                with st.form("login_form", clear_on_submit=False):
                    email = st.text_input("Email", placeholder="name@example.com", key="login_email")
                    password = st.text_input("Пароль", type="password", placeholder="Минимум 6 символов", key="login_password")
                    submit = st.form_submit_button("Войти", use_container_width=True, type="primary")

                    if submit:
                        if not email or not password:
                            st.error("Заполните все поля")
                        else:
                            try:
                                data = supabase_auth.auth.sign_in_with_password({
                                    "email": email,
                                    "password": password
                                })
                                st.session_state.user = data.user
                                # Сохраняем данные сессии для JS моста
                                st.session_state.session_data = {
                                    "access_token": data.session.access_token,
                                    "refresh_token": data.session.refresh_token,
                                    "expires_at": data.session.expires_at
                                }
                                st.rerun()
                            except Exception as e:
                                error_msg = str(e)
                                if "Invalid login credentials" in error_msg:
                                    st.error("❌ Неверный email или пароль")
                                elif "Email not confirmed" in error_msg:
                                    st.error("📧 Email не подтвержден. Проверьте почту.")
                                else:
                                    st.error(f"Ошибка входа: {error_msg}")

            with tab_register:
                with st.form("register_form", clear_on_submit=False):
                    new_email = st.text_input("Email", placeholder="name@example.com", key="reg_email")
                    new_password = st.text_input("Пароль", type="password", placeholder="Минимум 6 символов", key="reg_password")
                    new_password2 = st.text_input("Повторите пароль", type="password", placeholder="Минимум 6 символов", key="reg_password2")
                    submit_reg = st.form_submit_button("Создать аккаунт", use_container_width=True, type="primary")

                    if submit_reg:
                        if not new_email or not new_password or not new_password2:
                            st.error("Заполните все поля")
                        elif new_password != new_password2:
                            st.error("❌ Пароли не совпадают")
                        elif len(new_password) < 6:
                            st.error("❌ Пароль должен быть не менее 6 символов")
                        else:
                            try:
                                data = supabase_auth.auth.sign_up({
                                    "email": new_email,
                                    "password": new_password
                                })
                                if data.user and data.user.aud == "authenticated":
                                    st.session_state.user = data.user
                                    # Сохраняем данные сессии для JS моста
                                    if hasattr(data, 'session') and data.session:
                                        st.session_state.session_data = {
                                            "access_token": data.session.access_token,
                                            "refresh_token": data.session.refresh_token,
                                            "expires_at": data.session.expires_at
                                        }
                                    st.success("✅ Регистрация прошла успешно!")
                                    st.rerun()
                                else:
                                    st.success("✅ Аккаунт создан! Проверьте почту для подтверждения.")
                            except Exception as e:
                                error_msg = str(e)
                                if "User already registered" in error_msg:
                                    st.error("❌ Пользователь с таким email уже существует")
                                elif "rate limit" in error_msg.lower():
                                    st.error("⏳ Слишком много попыток. Подождите минуту.")
                                else:
                                    st.error(f"Ошибка регистрации: {error_msg}")

            st.markdown("<p style='text-align: center; color: var(--secondary-text); font-size: 0.75rem; margin-top: 10px;'>🔒 Данные защищены шифрованием</p>", unsafe_allow_html=True)

st.markdown(f"<p style='text-align: center; color: var(--secondary-text); font-weight: 500;'>Профессиональный юридический аудит договоров</p>", unsafe_allow_html=True)

# --- ОБНОВЛЕННЫЕ ТАРИФЫ (FUTURE LOGIC) ---
checkout_url = "https://jurisclearai.lemonsqueezy.com/checkout/buy/69a180c9-d5f5-4018-9dbe-b8ac64e4ced8"

st.markdown(f"""
<div class="pricing-card-container">
<!-- КАРТОЧКА 1: РАЗОВЫЙ АУДИТ -->
<div class="pricing-card">
<div class="pricing-header">Базовый</div>
<div class="pricing-price">850 ₽</div>
<ul class="pricing-features">
<li><b>Бесплатное резюме</b> основных рисков</li>
<li>Детальный юридический разбор (Full Report)</li>
<li>Конкретные правки для защиты интересов</li>
<li>Экспорт отчета в PDF и Word</li>
</ul>
<div style="background: rgba(59, 130, 246, 0.15); padding: 12px; border-radius: 12px; text-align: center; font-size: 11px; color: var(--accent-blue); margin-bottom: 20px;">
ℹ️ Оплачивайте только если результат вас устроит
</div>
<div style="width: 100%; background: rgba(255,255,255,0.05); color: var(--secondary-text); border: 1px dashed var(--border-color); padding: 14px; border-radius: 14px; text-align: center; font-weight: 600;">
Анализ доступен ниже 👇
</div>
</div>
<!-- КАРТОЧКА 2: БЕЗЛИМИТ PRO -->
<div class="pricing-card pricing-card-pro">
<div class="pricing-header" style="color: var(--accent-green);">Премиум</div>
<div class="pricing-price">2500 ₽ <span style="font-size: 14px; opacity: 0.6; font-weight: 400;">/мес</span></div>
<ul class="pricing-features">
<li><b>Неограниченное</b> количество документов</li>
<li>Полные отчеты <b>мгновенно</b> без доплат</li>
<li>Доступ к результату в истории навсегда</li>
<li>Персональный архив всех проверок</li>
<li>Самая мощная модель ИИ (GPT-4o)</li>
</ul>
<a href="{checkout_url}" target="_blank" style="display: block; background: linear-gradient(90deg, var(--accent-blue), var(--accent-cyan)); color: white; text-align: center; padding: 14px; border-radius: 14px; text-decoration: none; font-weight: 700; font-size: 15px; box-shadow: 0 4px 15px rgba(59, 130, 246, 0.3);">
🚀 Оформить подписку
</a>
</div>
</div>
""", unsafe_allow_html=True)

st.divider()

# Параметры анализа
st.markdown("### ⚙️ Параметры анализа")
c1, c2 = st.columns(2)

with c1:
    st.write("**Ваша роль:**")
    user_role = st.pills(
        "Роль", 
        [
            "Заказчик", "Исполнитель", 
            "Покупатель", "Поставщик", 
            "Арендатор", "Арендодатель", 
            "Работник", "Работодатель", 
            "Инвестор", "Основатель",
            "Лицензиат", "Лицензиар"
        ], 
        selection_mode="single", 
        default="Заказчик",
        label_visibility="collapsed",
        key=f"role_pills_{st.session_state.reset_counter}"
    )

with c2:
    st.write("**Тип документа:**")
    contract_type = st.pills(
        "Тип", 
        [
            "Авто-определение", "Услуги", 
            "Поставка / Купля-продажа", "NDA", 
            "Аренда", "Трудовой", 
            "ИТ-разработка", "Лицензионный", 
            "Займ", "Агентский"
        ], 
        selection_mode="single", 
        default="Авто-определение",
        label_visibility="collapsed",
        key=f"type_pills_{st.session_state.reset_counter}"
    )

# Рабочее пространство (Вкладки)
tab_audit, tab_redline, tab_demo = st.tabs(["🚀 ИИ Аудит", "🔄 Сравнение версий", "📝 Пример отчета"])

with tab_audit:
    # --- ЮРИДИЧЕСКИЙ ДИСКЛЕЙМЕР ---
    st.markdown("""
        <div style="background-color: #ff4b4b22; border: 2px solid #ff4b4b; padding: 15px; border-radius: 10px; margin-bottom: 20px;">
            <h4 style="margin-top: 0; color: #ff4b4b;">⚖️ Внимание: Юридический отказ от ответственности</h4>
            <p style="font-size: 0.9em; line-height: 1.4; margin-bottom: 0;">
                Данный сервис работает на базе искусственного интеллекта и <b>не является юридической консультацией</b>. 
                ИИ может ошибаться, галлюцинировать или пропускать важные детали. 
                Результаты анализа носят ознакоительный характер. Перед принятием решений обязательно 
                <b>проконсультируйтесь с квалифицированным юристом</b>. 
                Мы не несем ответственности за последствия использования данного инструмента.
            </p>
        </div>
    """, unsafe_allow_html=True)
    
    file = st.file_uploader("Выберите файл договора (PDF)", type=['pdf'], key=f"uploader_{st.session_state.reset_counter}")
    if file:
        if "analysis_result" not in st.session_state:
            if st.button("Начать анализ", use_container_width=True, type="primary"):
                with st.spinner("ИИ проводит глубокий юридический аудит..."):
                    try:
                        # --- ИНТЕГРИРОВАННЫЙ КУСОЧЕК КОДА ---
                        file_bytes = file.read() # Считываем байты один раз
                        text = extract_text_from_pdf(file_bytes)

                        if not text.strip():
                            st.error("Не удалось извлечь текст из документа. Возможно, файл поврежден или пуст.")
                            st.stop()
                        # -----------------------------------
                        
                    except Exception as e:
                        st.error(f"Ошибка при чтении PDF: {e}")
                        st.stop()
                    
                    special_instructions = ""
                    if contract_type == "NDA":
                        special_instructions = "Фокус на сроках конфиденциальности, исключениях и штрафaх за разглашение."
                    elif contract_type == "Аренда":
                        special_instructions = "Фокус на индексации цены, условиях расторжения, возврате депозита и текущем ремонте."
                    elif contract_type == "Трудовой":
                        special_instructions = "Фокус на дисциплинарных взысканиях, условиях увольнения, обязанностях и мат. ответственности."
                    elif contract_type == "ИТ-разработка":
                        special_instructions = "Фокус на передаче исключительных прав на код, этапах приемки и гарантийном периоде."
                    elif contract_type == "Поставка / Купля-продажа":
                        special_instructions = "Фокус на переходах рисков, сроках поставки, штрафах за недопоставку и скрытых дефектах."
                    elif contract_type == "Займ":
                        special_instructions = "Фокус на процентах, очередности погашения, штрафах за просрочку и условиях досрочного возврата."
                    elif contract_type == "Лицензионный":
                        special_instructions = "Фокус на территории использования, объеме прав, сублицензировании и роялти."
                    elif contract_type == "Агентский":
                        special_instructions = "Фокус на порядке отчетности агента, расчете вознаграждеия и праве на прямой поиск клиентов."

                    prompt_instruction = (
                        "СТРУКТУРА ОТВЕТА (ОБЯЗАТЕЛЬНО):\n"
                        "## ⚖️ Юридический анализ рисков\n"
                        "1. ОБЩИЙ ВЕРДИКТ.\n"
                        "2. ФИНАНСОВЫЕ РИСКИ.\n"
                        "3. РИСКИ РАСТОРЖЕНИЯ И СПОРОВ.\n"
                        "Для каждого риска пиши: 'Суть условия' и 'Юридический анализ'. Используй 🔴, 🟠, 🟡.\n\n"
                        "Будь строгим критиком. Если в договоре есть штрафы без вины или односторонние кабальные условия, "
                        "оценка риска (SCORE) должна быть высокой (7-10). "
                        "Разделяй пункты отчета двойным переносом строки для четкой читаемости.\n\n"
                        f"Действуй как опытный корпоративный юрист. Специализация: {contract_type}. "
                        f"Твоя задача — защитить интересы стороны: {user_role}. {special_instructions}\n\n"
                        "ЭТАЛОН КАЧЕСТВА АНАЛИЗА:\n"
                        "🔴 Критический риск: Несоразмерная неустойка (п. 6.2). Установлен штраф 1% в день. "
                        "Юридический анализ: Это 365% годовых, что в 10 раз выше рыночной нормы (0.1%).\n\n"
                        "ИНСТРУКЦИЯ ДВОЙНОЙ ПРОВЕРКИ (Chain of Verification):\n"
                        "Шаг 1: Проанализируй текст и выдели риски.\n"
                        "Шаг 2: Для каждого риска проверь, действительно ли в тексте договора есть указанный пункт и условие.\n"
                        "Шаг 3: Сформируй итоговый отчет. Если риск не подтвержден фактами — удали его.\n\n"
                        " ## 🛠️ Протокол разногласий (Готовые правки)\n"
                        "Составь таблицу в формате Markdown для всех найденных рисков:\n"
                        "| № Пункта | Оригинальный текст | Предлагаемая редакция | Обоснование |\n"
                        "| :--- | :--- | :--- | :--- |\n\n"
                        "В самом конце напиши 'SCORE: X' (1-10).\n"
                        "Язык: Русский."
                    )

                    # Используем функцию для анализа длинного текста
                    raw_res = analyze_long_text(text, contract_type, user_role, special_instructions, prompt_instruction)
                    
                    # Извлекаем оценку и очищаем текст для интерфейса
                    score_match = re.search(r"SCORE:\s*(\d+)", raw_res)
                    score = int(score_match.group(1)) if score_match else 5
                    clean_res = re.sub(r"SCORE:\s*\d+", "", raw_res).strip()

                    if clean_res:
                        st.session_state.analysis_result = clean_res
                        st.session_state.audit_score = score
                        st.rerun()
        else:
            # --- ИНТЕГРИРОВАННЫЙ БЛОК ВЫВОДА ОТЧЕТА ---
            score = st.session_state.get("audit_score", 5)
            bar_color, bar_shadow, risk_text = get_risk_params(score)
            st.write("### ИИ Оценка Риска:")
            st.markdown(f"""
                <div class="risk-meter-container">
                    <div class="risk-meter-bar" style="width:{score*10}%; background:{bar_color}; box-shadow: 0 4px 15px {bar_shadow};">
                        {risk_text} ({score}/10)
                    </div>
                </div>
            """, unsafe_allow_html=True)

            if "analysis_result" in st.session_state:
                st.success("✅ Анализ и протокол разногласий успешно сформированы!")

                clean_res = st.session_state.analysis_result
                
                # Показываем результат сразу
                st.markdown(f"<div class='report-card'>{clean_res.strip()}</div>", unsafe_allow_html=True)
                
                # Три колонки для кнопок (ID заменен на фейковый или удален)
                col_pdf, col_word, col_sup = st.columns(3)
                
                with col_pdf:
                    pdf_bytes = create_pdf(clean_res)
                    if pdf_bytes:
                        st.download_button(
                            label="📥 PDF",
                            data=bytes(pdf_bytes),
                            file_name=f"audit_report.pdf",
                            mime="application/pdf",
                            use_container_width=True
                        )
                    else:
                        st.warning("PDF не доступен")
                
                with col_word:
                    try:
                        word_bytes = create_docx(clean_res)
                        if word_bytes:
                            st.download_button(
                                label="📝 Word",
                                data=word_bytes,
                                file_name=f"audit_report.docx",
                                mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
                                use_container_width=True
                            )
                        else:
                            st.warning("Word не доступен")
                    except Exception as e:
                        st.error("Ошибка Word")
                
                with col_sup:
                    st.link_button("🆘 Поддержка", "https://t.me/твой_логин", use_container_width=True)

                st.write("")
                if st.button("📁 Загрузить новый договор", use_container_width=True, key="btn_paid_reset"):
                    st.session_state.reset_counter += 1
                    keys_to_clear = ["analysis_result", "audit_score"]
                    for k in keys_to_clear:
                        if k in st.session_state: del st.session_state[k]
                    st.rerun()
with tab_demo:
    st.write("### Так выглядит результат анализа:")
    bar_color, bar_shadow, risk_text = get_risk_params(9)
    st.markdown(f"""
        <div class="risk-meter-container">
            <div style="height:35px; width:90%; background:{bar_color}; 
            box-shadow: 0 4px 15px {bar_shadow}; border-radius:10px; 
            display:flex; align-items:center; justify-content:center; color:white; font-weight:900;">
                КРИТИЧЕСКИЙ (9/10)
            </div>
        </div>
    """, unsafe_allow_html=True)
    st.markdown(f"<div class='report-card'>{sample_text}</div>", unsafe_allow_html=True)

st.divider()
col_f1, col_f2, col_f3 = st.columns(3)
with col_f1:
    st.caption("© 2026 JurisClear AI | Ереван")
with col_f2:
    if st.button("Конфиденциальность", type="tertiary"):
        st.info("Политика конфиденциальности...")
with col_f3:
    if st.button("Условия", type="tertiary"):
        st.info("Условия использования...")

st.markdown(f"<p style='text-align: center; color: var(--secondary-text); font-size: 11px; margin-top: 20px;'>support@jurisclear.com</p>", unsafe_allow_html=True)
